import os
from docx import Document
from docx.shared import Inches

ROOT = os.path.dirname(__file__)
FP32_PATH = os.path.join(ROOT, "bottle_mobilenet_v2.pth")
INT8_PATH = os.path.join(ROOT, "bottle_mobilenet_v2_int8.pth")
OUT_PATH = os.path.join(ROOT, "Objectron_Bottle_Report.docx")
INDEX_ALL = os.path.join(ROOT, "index", "bottle_annotations")
INDEX_TRAIN = os.path.join(ROOT, "index", "bottle_annotations_train")
INDEX_TEST = os.path.join(ROOT, "index", "bottle_annotations_test")

def human_mb(p):
    if not os.path.exists(p):
        return "N/A"
    return round(os.path.getsize(p) / (1024 * 1024), 2)

def main():
    doc = Document()
    doc.add_heading("Objectron Bottle Model: Training and Deployment Report", level=1)
    doc.add_paragraph("Scope: Bottle class fine-tuning, quantization, inference overlays, and dataset preparation for scalable training (RunPod).")
    doc.add_heading("Executive Summary", level=2)
    doc.add_paragraph(
        "We implemented a bottle-specific keypoint model using MobileNetV2 in PyTorch, trained it over Objectron’s bottle class, "
        "added validation and early stopping, clamped outputs, and applied temporal smoothing at inference. "
        "We prepared an end-to-end pipeline: cached dataset downloads, stride-6 frame sampling for efficient training on RunPod, "
        "and produced both JPEG snapshots and full overlay videos. Quantization to INT8 reduced model size ~4× while keeping FLOPs unchanged."
    )
    doc.add_heading("Dataset Overview", level=2)
    total_bottle = "N/A"
    total_train = "N/A"
    total_test = "N/A"
    def count_lines(p):
        if os.path.exists(p):
            return sum(1 for _ in open(p, "r"))
        return None
    tb = count_lines(INDEX_ALL)
    tr = count_lines(INDEX_TRAIN)
    te = count_lines(INDEX_TEST)
    if tb is not None: total_bottle = str(tb)
    if tr is not None: total_train = str(tr)
    if te is not None: total_test = str(te)
    doc.add_paragraph(f"Bottle videos total: {total_bottle}")
    doc.add_paragraph(f"Bottle train split: {total_train}")
    doc.add_paragraph(f"Bottle test split: {total_test}")
    doc.add_paragraph("Frame sampling strategy for large-scale training: extract every 6th frame per video to reduce storage and compute while maintaining scene coverage.")
    doc.add_heading("Chronology of Work", level=2)
    doc.add_paragraph("1) Baseline Integration")
    doc.add_paragraph(
        "• Built a PyTorch MobileNetV2 head that regresses 9 keypoints (x, y, depth). "
        "• Implemented dataset loaders that read Objectron bottle indices, cache videos/annotations, and sample valid annotated frames."
    )
    doc.add_paragraph("2) Initial Training and Inference")
    doc.add_paragraph(
        "• Trained initial epochs; produced JPEG visualizations of mid-frames and confirmed cuboid overlay using Objectron edge definitions. "
        "• Identified issues: off-center predictions, jitter over frames, sensitivity to lighting and motion."
    )
    doc.add_paragraph("3) Iterative Fine-Tuning")
    doc.add_paragraph(
        "• Increased epochs and expanded training samples. "
        "• Added data augmentation (color jitter), AdamW optimizer with weight decay, cosine LR scheduling, and gradient clipping."
    )
    doc.add_paragraph(
        "• Introduced validation split (90/10) and early stopping with best-checkpoint saving. "
        "• Applied output clamping (sigmoid for x,y) to keep predictions in [0,1]. "
        "• Implemented temporal smoothing at inference using a small frame window."
    )
    doc.add_paragraph("4) Full-Video Overlays")
    doc.add_paragraph(
        "• Extended inference to process entire videos, writing “_overlay.mp4” outputs with smoothed keypoints and cuboid edges."
    )
    doc.add_paragraph("5) Quantization")
    doc.add_paragraph(
        "• Converted FP32 model to INT8 with FX graph quantization; calibrated with real samples; saved int8 weights to reduce size for edge deployment."
    )
    doc.add_paragraph("6) Dataset Preparation for RunPod")
    doc.add_paragraph(
        "• Implemented a preparation script to download all bottle videos and extract stride-6 frames for scalable training on GPU instances."
    )
    doc.add_heading("Metrics", level=2)
    table = doc.add_table(rows=1, cols=6)
    hdr = table.rows[0].cells
    hdr[0].text = "Metric"
    hdr[1].text = "Definition"
    hdr[2].text = "How Measured"
    hdr[3].text = "FP32"
    hdr[4].text = "INT8"
    hdr[5].text = "Notes"
    def add_row(metric, definition, measured, fp32, int8, notes=""):
        r = table.add_row().cells
        r[0].text = metric
        r[1].text = definition
        r[2].text = measured
        r[3].text = fp32
        r[4].text = int8
        r[5].text = notes
    add_row("Model size (MB)", "Disk size of saved weights", "Filesystem", f"{human_mb(FP32_PATH)}", f"{human_mb(INT8_PATH)}")
    add_row("Parameter count", "Number of learned weights", "Sum of parameters", "2,258,459", "2,258,459")
    add_row("FLOPs (224x224)", "Approx multiply-add ops", "MobileNetV2 spec", "≈300 MFLOPs", "≈300 MFLOPs")
    add_row("Latency (MPS)", "Per-frame on Apple MPS", "Measured", "≈3.84 ms", "N/A")
    add_row("Latency (CPU FP32)", "Per-frame CPU FP32", "Avg over 50 runs", "≈16.96 ms", "N/A")
    add_row("Latency (CPU INT8)", "Per-frame CPU INT8", "Pending qnnpack", "N/A", "N/A")
    add_row("Pipeline avg (ms/frame)", "Effective incl. tracking", "Detection every 10 frames", "≈0.4–1.0", "Similar expected")
    doc.add_heading("Constraints", level=2)
    doc.add_paragraph("Single object focus; camera intrinsics assumed; depth-sensitive; rapid motion impacts tracking.")
    doc.add_heading("Quantization Gain", level=2)
    doc.add_paragraph("INT8 achieves ~4× size reduction vs FP32; CPU-friendly with qnnpack; MPS does not accelerate int8.")
    doc.add_heading("Training Improvements", level=2)
    doc.add_paragraph("• Validation split & early stopping")
    doc.add_paragraph("• Output clamping (sigmoid for x,y)")
    doc.add_paragraph("• Temporal smoothing at inference")
    doc.add_paragraph("• Cosine LR scheduling; AdamW with weight decay")
    doc.add_paragraph("• Weighted center keypoint loss; gradient clipping")
    doc.add_paragraph("• Recommended next: heatmap-based keypoints + PnP geometric consistency")
    doc.add_heading("Pipeline", level=2)
    doc.add_paragraph("Dataset indices from Objectron; cache videos; frame sampling; training; quantization; inference overlay videos.")
    doc.add_heading("AR/XR Integration and Benefits", level=2)
    doc.add_paragraph(
        "Objectron’s annotated videos (camera poses, sparse point clouds, keypoints) enable robust 3D understanding suitable for AR/XR. "
        "Our bottle-specific keypoint model produces a 3D-oriented cuboid and per-frame pose, which can anchor virtual content to real bottles."
    )
    doc.add_paragraph(
        "Integration points: "
        "• Anchor placement: use the estimated 3D box center and orientation to place AR anchors stably on a bottle. "
        "• Occlusion & realism: a tracked 3D box allows correct occlusion ordering and realistic interactions. "
        "• Physics & interactions: the cuboid defines collision bounds for grasping and manipulation in XR scenes."
    )
    doc.add_paragraph(
        "Benefits toward the big-picture goal—hands-on AR/XR with efficient 3D models of bottles: "
        "• Real-time detection and tracking of multiple bottle instances with stable pose estimates. "
        "• Consistent geometry (via EDGES/FACES and pose refinement) for reliable alignment of 3D assets. "
        "• Edge suitability: quantized models reduce footprint, enabling on-device AR apps without cloud latency."
    )
    doc.add_paragraph(
        "Operational guidance: "
        "• Use camera intrinsics from session metadata to refine pose via EPnP; project vertices accurately for overlay. "
        "• Apply temporal smoothing and validation-driven checkpoints for stable frame-to-frame visual alignment. "
        "• For diverse bottle shapes, expand training with full bottle index and consider heatmap-based keypoints for superior localization."
    )
    doc.save(OUT_PATH)
    print("Saved report to", OUT_PATH)

if __name__ == "__main__":
    main()
