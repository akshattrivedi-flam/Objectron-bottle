import os
import sys
import numpy as np
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from google.protobuf.json_format import MessageToDict

# Add the directory containing the schema to sys.path
sys.path.append(os.path.join(os.getcwd(), 'Objectron'))

try:
    from objectron.schema import annotation_data_pb2
except ImportError:
    sys.path.append(os.getcwd())
    from objectron.schema import annotation_data_pb2

def create_detailed_analysis_doc():
    pb_file = "Objectron/bottle_annotation.pbdata"
    img_path = "Objectron/bottle_result.jpg"
    
    if not os.path.exists(pb_file):
        print(f"File not found: {pb_file}")
        return

    # 1. Extract data from Proto
    sequence = annotation_data_pb2.Sequence()
    with open(pb_file, 'rb') as f:
        sequence.ParseFromString(f.read())
    
    data = MessageToDict(sequence)
    obj = data['objects'][0]
    frame = data['frameAnnotations'][0]
    
    # 2. Setup Document
    doc = Document()
    
    title = doc.add_heading('Deep Dive: Objectron 3D Annotation Analysis', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # SECTION 1: VISUAL REFERENCE
    doc.add_heading('1. Visual Reference: Real-World Sample', level=1)
    doc.add_paragraph(
        "The following image represents a typical real-world capture from the Objectron bottle dataset. "
        "This is the raw input before any geometric processing or annotation is applied."
    )
    
    # Add the provided image if it exists
    img_path_sample = "Objectron/images/bottle_sample.jpg"
    if os.path.exists(img_path_sample):
        doc.add_picture(img_path_sample, width=Inches(4))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph("[Image Placeholder: bottle_sample.jpg not found]")

    # SECTION 2: THE ANNOTATION DATA (FROM PBDATA)
    doc.add_heading('2. The Annotation Data (From .pbdata)', level=1)
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(4.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("The image above shows the final 3D bounding box projected onto the 2D plane using the matrices decoded below.")

    # SECTION 3: MATRIX BREAKDOWN
    doc.add_heading('3. Mathematical Matrix Breakdown (Exhaustive)', level=1)
    doc.add_paragraph("This section defines every row and column of the matrices found in the .pbdata file, explaining their physical and geometric roles.")

    # A. Object Pose in World Coordinate System
    doc.add_heading('A. Object Model Matrix (Object-to-World)', level=2)
    doc.add_paragraph("This matrix places the bottle in the 3D room. It consists of Scale, Rotation, and Translation.")
    
    # Scale
    scale = obj['scale']
    doc.add_heading('1. Scale (3x1 Vector)', level=3)
    doc.add_paragraph(f"Defines the size of the bottle in meters: Width={scale[0]:.4f}m, Height={scale[1]:.4f}m, Depth={scale[2]:.4f}m.")
    
    # Rotation
    rot = np.array(obj['rotation']).reshape(3, 3)
    doc.add_heading('2. Object Rotation (3x3 Matrix)', level=3)
    doc.add_paragraph("Defines how the bottle is oriented in the world.")
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    for i in range(3):
        for j in range(3):
            table.cell(i, j).text = f"{rot[i, j]:.4f}"
    
    bullets = [
        "Column 1 (X-axis): The 'Right' direction of the bottle.",
        "Column 2 (Y-axis): The 'Up' direction of the bottle.",
        "Column 3 (Z-axis): The 'Forward' direction of the bottle."
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    # Translation
    trans = obj['translation']
    doc.add_heading('3. Object Translation (3x1 Vector)', level=3)
    doc.add_paragraph(f"The 3D coordinates (X, Y, Z) of the bottle's center in the world: ({trans[0]:.4f}, {trans[1]:.4f}, {trans[2]:.4f}).")

    # B. Camera View Matrix (World-to-Camera)
    doc.add_heading('B. Camera View & Transform (Extrinsics)', level=2)
    
    # 1. Transform (Camera-to-World)
    transform = np.array(frame['camera']['transform']).reshape(4, 4)
    doc.add_heading('1. Camera Transform Matrix (Camera-to-World)', level=3)
    doc.add_paragraph("This matrix describes where the camera is located in the world coordinate system.")
    
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'
    for i in range(4):
        for j in range(4):
            table.cell(i, j).text = f"{transform[i, j]:.4f}"

    doc.add_paragraph("Physical Interpretation:")
    bullets = [
        "Column 4 (Top 3): The (X, Y, Z) position of the phone in the room.",
        "Upper 3x3: The orientation of the phone (where it's pointing)."
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    # 2. View Matrix (World-to-Camera)
    view = np.array(frame['camera']['viewMatrix']).reshape(4, 4)
    doc.add_heading('2. Camera View Matrix (World-to-Camera)', level=3)
    doc.add_paragraph("The inverse of the Transform. It shifts the world so that the camera is at the origin (0,0,0).")
    
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'
    for i in range(4):
        for j in range(4):
            table.cell(i, j).text = f"{view[i, j]:.4f}"
    
    doc.add_paragraph("Calculation: View = Transform^-1")

    # 3. Euler Angles
    angles = frame['camera']['eulerAngles']
    doc.add_heading('3. Camera Orientation (Euler Angles)', level=3)
    doc.add_paragraph(f"The tilt of the phone in radians: Roll={angles['roll']:.4f}, Pitch={angles['pitch']:.4f}, Yaw={angles['yaw']:.4f}.")

    # C. Camera Intrinsic Matrix (Camera-to-Image)
    doc.add_heading('C. Camera Intrinsic Matrix (K)', level=2)
    intrinsics = np.array(frame['camera']['intrinsics']).reshape(3, 3)
    doc.add_paragraph("This 3x3 matrix defines the lens properties of the phone camera.")
    
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    for i in range(3):
        for j in range(3):
            table.cell(i, j).text = f"{intrinsics[i, j]:.4f}"

    bullets = [
        f"fx (Row 1, Col 1): {intrinsics[0,0]:.2f} - Horizontal Focal Length.",
        f"fy (Row 2, Col 2): {intrinsics[1,1]:.2f} - Vertical Focal Length.",
        f"cx (Row 1, Col 3): {intrinsics[0,2]:.2f} - Principal Point X (Optical Center).",
        f"cy (Row 2, Col 3): {intrinsics[1,2]:.2f} - Principal Point Y (Optical Center)."
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    # D. Projection Matrix (Clip Space)
    doc.add_heading('D. Projection Matrix', level=2)
    proj = np.array(frame['camera']['projectionMatrix']).reshape(4, 4)
    doc.add_paragraph("This 4x4 matrix maps 3D points into the normalized 'Clip Space' for rendering.")
    
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'
    for i in range(4):
        for j in range(4):
            table.cell(i, j).text = f"{proj[i, j]:.4f}"

    doc.add_paragraph("Key Elements:")
    bullets = [
        "Row 1, Col 1: Controls the Horizontal Field of View.",
        "Row 2, Col 2: Controls the Vertical Field of View.",
        "Row 3, Col 3 & 4: Define the Near and Far clipping planes (what's too close or too far to see)."
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    # E. Surface Plane Data
    doc.add_heading('E. Ground Truth Surface Plane', level=2)
    plane_center = frame.get('planeCenter', [0,0,0])
    plane_normal = frame.get('planeNormal', [0,0,0])
    doc.add_paragraph(
        "Objectron detects the flat surface (e.g., table, floor) the object is sitting on. "
        "This is used to ensure the 3D box doesn't 'float' in the air."
    )
    doc.add_paragraph(f"Plane Center (X, Y, Z): ({plane_center[0]:.4f}, {plane_center[1]:.4f}, {plane_center[2]:.4f})")
    doc.add_paragraph(f"Plane Normal (Direction Up): ({plane_normal[0]:.4f}, {plane_normal[1]:.4f}, {plane_normal[2]:.4f})")

    # SECTION 4: THE TARGET KEYPOINTS
    doc.add_heading('4. Target 2D Keypoints (Final Labels)', level=1)
    doc.add_paragraph("These are the normalized [0, 1] coordinates the MobileNetV2 model is trained to predict.")
    
    kp_data = frame['annotations'][0]['keypoints']
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Keypoint ID'
    hdr_cells[1].text = 'X (Normalized)'
    hdr_cells[2].text = 'Y (Normalized)'
    
    for kp in kp_data:
        row_cells = table.add_row().cells
        row_cells[0].text = str(kp.get('id', 0))
        row_cells[1].text = f"{kp['point2d']['x']:.4f}"
        row_cells[2].text = f"{kp['point2d']['y']:.4f}"

    # SECTION 5: 9-KEYPOINT CALCULATION (THE "HOW")
    doc.add_heading('5. The 9-Keypoint Calculation (The "How")', level=1)
    doc.add_paragraph(
        "This section explains exactly how the 9 points are derived from the Scale (S), Rotation (R), and Translation (T) "
        "stored in the .pbdata file."
    )

    doc.add_heading('Step 1: The Canonical Template', level=2)
    doc.add_paragraph(
        "Every object starts as a 'Unit Box' template. The coordinates are predefined as follows:"
    )
    
    template_table = doc.add_table(rows=1, cols=4)
    template_table.style = 'Table Grid'
    hdr = template_table.rows[0].cells
    hdr[0].text = 'ID'
    hdr[1].text = 'Name'
    hdr[2].text = 'Template (X, Y, Z)'
    hdr[3].text = 'Formula (Local)'
    
    template_points = [
        (0, "Centroid", "(0, 0, 0)", "Origin"),
        (1, "Corner 1", "(-0.5, -0.5, -0.5)", "(-w/2, -h/2, -d/2)"),
        (2, "Corner 2", "(-0.5, -0.5, 0.5)", "(-w/2, -h/2, +d/2)"),
        (3, "Corner 3", "(-0.5, 0.5, -0.5)", "(-w/2, +h/2, -d/2)"),
        (4, "Corner 4", "(-0.5, 0.5, 0.5)", "(-w/2, +h/2, +d/2)"),
        (5, "Corner 5", "(0.5, -0.5, -0.5)", "(+w/2, -h/2, -d/2)"),
        (6, "Corner 6", "(0.5, -0.5, 0.5)", "(+w/2, -h/2, +d/2)"),
        (7, "Corner 7", "(0.5, 0.5, -0.5)", "(+w/2, +h/2, -d/2)"),
        (8, "Corner 8", "(0.5, 0.5, 0.5)", "(+w/2, +h/2, +d/2)")
    ]
    
    for pid, name, coords, formula in template_points:
        row = template_table.add_row().cells
        row[0].text = str(pid)
        row[1].text = name
        row[2].text = coords
        row[3].text = formula

    doc.add_heading('Step 2: Transforming to World Space', level=2)
    doc.add_paragraph(
        "To get the final 3D points in the world, we apply the following calculation for EACH of the 9 points:"
    )
    doc.add_paragraph("P_world = (Rotation * (Scale * P_template)) + Translation")
    
    doc.add_paragraph("Example Calculation for Centroid (Point 0):")
    doc.add_paragraph(
        "P_world(0) = (Rotation * (Scale * [0,0,0])) + Translation\n"
        "P_world(0) = (Rotation * [0,0,0]) + Translation\n"
        "P_world(0) = [0,0,0] + Translation\n"
        "Therefore, Keypoint 0 is exactly the Translation vector."
    ).paragraph_format.left_indent = Inches(0.3)

    doc.add_paragraph("Example Calculation for Corner 1:")
    doc.add_paragraph(
        "P_world(1) = (Rotation * ([w,h,d] * [-0.5, -0.5, -0.5])) + Translation\n"
        "P_world(1) = (Rotation * [-w/2, -h/2, -d/2]) + Translation"
    ).paragraph_format.left_indent = Inches(0.3)

    doc.add_heading('Step 3: The Full 3D-to-2D Projection Pipeline', level=2)
    doc.add_paragraph(
        "The final 2D point (u, v) on your phone screen is calculated by chaining all the matrices discussed in Section 3:"
    )
    
    formula = "P_2d = K * [ViewMatrix] * [ModelMatrix] * P_template"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(formula)
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph("Breaking down the pipeline:")
    steps = [
        "1. Model Matrix: Places the object in the 3D room.",
        "2. View Matrix: Shifts the room so the camera is the center of the universe.",
        "3. Intrinsic Matrix (K): Projects the 3D 'Camera Space' point onto the 2D 'Image Plane'.",
        "4. Normalization: The result is divided by Depth (Z) to get final pixel coordinates."
    ]
    for s in steps:
        doc.add_paragraph(s, style='List Number')

    doc.save("Objectron/Objectron_Annotation_Analysis.docx")
    print("Analysis saved to Objectron/Objectron_Annotation_Analysis.docx")

if __name__ == "__main__":
    create_detailed_analysis_doc()
