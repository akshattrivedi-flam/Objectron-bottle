import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_kt_master_guide():
    doc = Document()
    
    # --- TITLE SECTION ---
    title = doc.add_heading('Knowledge Transfer (KT): Objectron for 3D Object Tracking', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # --- SECTION 1: INTRODUCTION & TRADITIONAL MATH ---
    doc.add_heading('1. The Evolution of 3D Object Tracking', level=1)
    
    doc.add_heading('Traditional Mathematical Approaches (Classical CV)', level=2)
    doc.add_paragraph(
        "Historically, 3D tracking relied on manually engineered features and complex geometric solvers. "
        "Techniques like SIFT, SURF, and ORB were used to find keypoints, which were then matched across frames."
    )
    
    doc.add_heading('The "Inefficiency" of Traditional Math:', level=3)
    bullets = [
        "Feature Sensitivity: Traditional math fails in low light, motion blur, or on textureless objects (like a plain glass bottle).",
        "Manual Modeling: You had to provide a perfect 3D CAD model of the object beforehand.",
        "Computational Heavy: Solving the PnP (Perspective-n-Point) problem for every frame using iterative optimization is CPU-intensive.",
        "Drift: Small errors in math accumulate, causing the 3D box to 'drift' away from the object over time."
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    doc.add_heading('The Objectron Advantage (Deep Learning)', level=2)
    doc.add_paragraph(
        "Objectron shifts the burden from 'Manual Math' to 'Learned Geometry'. Instead of us telling the computer how a bottle looks "
        "from every angle, we train a Deep Learning model (MobileNetV2) to 'see' the 3D structure."
    )
    bullets = [
        "Automation: The model automatically learns the most robust features for detection.",
        "Real-time Performance: By using MobileNetV2, we achieve high-speed tracking on mobile devices.",
        "Single-Image 3D: Unlike traditional SLAM, Objectron can estimate a 3D pose from just one single 2D frame."
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    # --- SECTION 2: OUR APPROACH - FINE-TUNING & OVERFITTING ---
    doc.add_heading('2. Our Approach: Targeted Fine-Tuning', level=1)
    doc.add_paragraph(
        "For AR/XR applications targeting specific branded products (like a specific bottle), a general model isn't enough. "
        "We use a strategy called 'Targeted Overfitting' or high-precision fine-tuning."
    )
    
    doc.add_heading('The Strategy:', level=2)
    steps = [
        "Data Specialization: We use the Objectron 'Bottle' subset (1,543 sequences) to teach the model the general category.",
        "Stride-8 Extraction: We extract frames every 8th step to ensure diverse viewpoints while keeping the dataset manageable.",
        "Overfitting for Precision: By training for 20 epochs without early stopping, we force the model to master the specific geometry of the target object class.",
        "AR/XR Integration: The resulting model provides the 9 keypoints needed to anchor digital content (AR) onto the physical bottle."
    ]
    for s in steps:
        doc.add_paragraph(s, style='List Number')

    # --- SECTION 3: THE MATHEMATICAL ENGINE ---
    doc.add_heading('3. The Mathematical Engine (The "How")', level=1)
    doc.add_paragraph("This is the core logic that powers the 3D-to-2D transformation.")

    doc.add_heading('The 9-Keypoint Proxy Model', level=2)
    doc.add_paragraph(
        "We don't predict a 3D box directly. We predict 9 points in 2D (Centroid + 8 Corners). "
        "This is the 'Geometric Proxy'."
    )
    
    doc.add_heading('The Projection Pipeline Formula:', level=3)
    formula = "P_2d = K * [ViewMatrix] * [ModelMatrix] * P_template"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(formula)
    run.bold = True
    run.font.size = Pt(12)

    doc.add_heading('Breakdown of Components:', level=3)
    components = [
        ("P_template:", "The 'Unit Cube' coordinates (Local Space)."),
        ("ModelMatrix (M):", "Combines Scale, Rotation, and Translation to place the bottle in the World."),
        ("ViewMatrix (V):", "Shifts the world so the camera is the center (Camera Space)."),
        ("Intrinsic Matrix (K):", "The lens properties that project 3D depth into 2D pixels.")
    ]
    for title, desc in components:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(title)
        run.bold = True
        p.add_run(f" {desc}")

    # --- NEW SUBSECTION: CAMERA POSE MATH ---
    doc.add_heading('Deep Dive: Camera Pose Calculation (VIO)', level=2)
    doc.add_paragraph(
        "The most critical component of the pipeline is the Camera Pose. In the Objectron dataset, "
        "this is calculated using Visual-Inertial Odometry (VIO), which fuses camera images with IMU sensors."
    )
    
    doc.add_heading('The Transform Matrix (T):', level=3)
    doc.add_paragraph(
        "Describes where the camera is in the world. It is a 4x4 matrix combining a 3x3 Rotation (R) and a 3x1 Translation (t)."
    )
    
    formula_t = "T = [[R, t], [0, 1]]"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(formula_t).bold = True

    doc.add_heading('The View Matrix (V) - The Inverse Math:', level=3)
    doc.add_paragraph(
        "To render the world FROM the camera's perspective, we must invert the Transform matrix. "
        "Because R is orthonormal, its inverse is its transpose (R^T), making the calculation efficient:"
    )
    
    formula_v = "V = T^-1 = [[R^T, -R^T * t], [0, 1]]"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(formula_v).bold = True

    doc.add_heading('VIO Fusion Strategy:', level=3)
    vio_bullets = [
        "IMU (Inertial): Measures high-speed motion but drifts over time.",
        "Visual (Features): Tracks pixels across frames to solve the Essential Matrix equation (x'^T E x = 0).",
        "Optimization: The system minimizes the Reprojection Error to provide a stable, drift-free pose for every frame."
    ]
    for b in vio_bullets:
        doc.add_paragraph(b, style='List Bullet')

    # --- SECTION 4: SUMMARY & TAKEAWAYS ---
    doc.add_heading('4. Key Takeaways for the Team', level=1)
    takeaways = [
        "Data is the New Math: We use AR-captured camera poses to generate ground truth, replacing manual annotation.",
        "Mobile-First: The architecture is optimized for low-latency AR/XR environments.",
        "Geometric Constraints: By predicting 9 related points, the model maintains 3D consistency even under occlusion."
    ]
    for t in takeaways:
        doc.add_paragraph(t, style='List Bullet')

    # --- SECTION 5: GAMMA.APP PPT PROMPT ---
    doc.add_heading('5. Gamma.app Presentation Prompt', level=1)
    doc.add_paragraph(
        "Copy and paste the following prompt into Gamma.app to generate a professional presentation for this KT session:"
    )
    
    prompt_box = doc.add_table(rows=1, cols=1)
    prompt_box.style = 'Table Grid'
    prompt_text = (
        "Create a professional, technical presentation about 'Objectron: Revolutionizing 3D Object Tracking for AR/XR'.\n\n"
        "Outline:\n"
        "1. Introduction: The shift from traditional manual math (SIFT/PnP) to Deep Learning (Objectron).\n"
        "2. The Problem: Why traditional math fails for textureless objects like bottles (Inefficiency, Drift, Manual Effort).\n"
        "3. The Solution: Objectron's learned geometry and the MobileNetV2 architecture.\n"
        "4. Technical Deep Dive: The 9-Keypoint model and the mathematical projection pipeline (Model, View, and Intrinsic matrices).\n"
        "5. Project Workflow: Training on 1543 bottle sequences, Stride-8 frame extraction, and targeted fine-tuning for branded products.\n"
        "6. Application in AR/XR: How 9 keypoints enable stable tracking and digital anchoring.\n\n"
        "Style: Modern, Tech-focused, using clean diagrams and geometric motifs."
    )
    prompt_box.cell(0, 0).text = prompt_text

    # Save
    save_path = 'Objectron/Objectron_KT_Master_Guide.docx'
    doc.save(save_path)
    print(f"KT Master Guide saved to {save_path}")

if __name__ == "__main__":
    create_kt_master_guide()
