from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_objectron_doc():
    doc = Document()
    
    # Title
    title = doc.add_heading('Objectron Dataset: Technical Architecture & Mathematical Interpretation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # SECTION: BIG PICTURE FLOW
    doc.add_heading('0. The Big Picture Flow (Data Journey)', level=1)
    doc.add_paragraph("The following flow chart illustrates the end-to-end journey of the Objectron data, from real-world capture to the final AI model prediction.")
    
    # Create the Flow Chart using text boxes and arrows
    flow_steps = [
        ("AR CAPTURE", "Phone sensors (ARCore/ARKit) record Video + Camera Pose + IMU data."),
        ("3D ANNOTATION", "Human annotators define a fixed 3D bounding box in World Space."),
        ("MATH PROJECTION", "Matrices (S, R, V, K) project 3D corners onto every 2D video frame."),
        ("DATA PREPARATION", "Frames are extracted (Stride-8) and paired with 9-keypoint labels."),
        ("MODEL TRAINING", "MobileNetV2 learns to predict 2D keypoints from raw RGB pixels."),
        ("3D INFERENCE", "Trained model predicts 3D pose from a single 2D image.")
    ]
    
    for i, (step, desc) in enumerate(flow_steps):
        # Step Box
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f" [ {step} ] ")
        run.bold = True
        run.font.size = Pt(14)
        
        # Description
        p_desc = doc.add_paragraph(desc)
        p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_desc.runs[0].font.size = Pt(10)
        
        # Arrow (except for the last step)
        if i < len(flow_steps) - 1:
            p_arrow = doc.add_paragraph("↓")
            p_arrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_arrow.runs[0].bold = True
            p_arrow.runs[0].font.size = Pt(16)

    # Introduction
    doc.add_heading('1. The AR-Capture Foundation', level=1)
    p = doc.add_paragraph(
        "The Objectron dataset is fundamentally different from traditional 2D computer vision datasets. "
        "It was captured using AR-capable mobile devices (ARCore/ARKit), which allow for the recording of "
        "camera poses and environment geometry alongside the video stream."
    )
    
    # Coordinate Systems
    doc.add_heading('2. Coordinate System Hierarchy', level=1)
    doc.add_paragraph("There are four primary coordinate systems involved in the data generation process:")
    
    systems = [
        ("Object Coordinate System:", "Canonical space where the object is at the center (0,0,0) and normalized to unit scale."),
        ("World Coordinate System:", "The 3D space of the room where the capture took place, established by the AR session."),
        ("Camera Coordinate System:", "A 3D space relative to the phone's camera lens (Origin is the optical center)."),
        ("Image Coordinate System:", "The 2D plane (pixels) of the recorded video frame.")
    ]
    
    for title, desc in systems:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(title)
        run.bold = True
        p.add_run(f" {desc}")

    # Mathematical Pipeline
    doc.add_heading('3. Detailed Mathematical Projection Pipeline', level=1)
    doc.add_paragraph(
        "To project a 3D point from the real world onto a 2D screen, we follow a rigorous linear algebra pipeline. "
        "Each step below represents a transformation from one coordinate space to the next."
    )
    
    # Step 1: Model Matrix
    doc.add_heading('Step 1: Local Space to World Space (The Model Matrix)', level=2)
    doc.add_paragraph(
        "Every object starts as a 'Unit Cube' in its own local coordinate system. To place it in the real world, "
        "we apply a Model Matrix (M), which is a combination of Scale (S), Rotation (R), and Translation (T)."
    )
    doc.add_paragraph("Formula: P_world = [R * S | t] * [P_local | 1]^T")
    doc.add_paragraph(
        "Derivation:\n"
        "1. Scaling: Each local point (x,y,z) is multiplied by the dimensions (w,h,d) to get the physical size.\n"
        "2. Rotation: The scaled points are rotated using a 3x3 orthonormal matrix to match the object's orientation.\n"
        "3. Translation: The object is shifted by vector 't' to its final position in the AR room."
    ).paragraph_format.left_indent = Inches(0.3)

    # Step 2: View Matrix
    doc.add_heading('Step 2: World Space to Camera Space (The View Matrix)', level=2)
    doc.add_paragraph(
        "The Camera (phone) is constantly moving. We must transform all world points into 'Camera Space', "
        "where the camera lens is at the origin (0,0,0) and looking down the Z-axis."
    )
    doc.add_paragraph("Formula: P_camera = V * P_world")
    doc.add_paragraph(
        "Derivation:\n"
        "The View Matrix (V) is the inverse of the Camera's Pose matrix. If the phone is at position 'C' with orientation 'Rc', "
        "then the transformation to bring the world TO the camera is:\n"
        "V = [Rc^T | -Rc^T * C]\n"
        "This effectively 'zeroes out' the camera's position, making all points relative to the lens."
    ).paragraph_format.left_indent = Inches(0.3)

    # Step 3: Projection & Intrinsic Matrix
    doc.add_heading('Step 3: Camera Space to Image Plane (The Intrinsic Matrix)', level=2)
    doc.add_paragraph(
        "This is the final step where 3D depth is 'squashed' into 2D pixels. This is defined by the camera's physical lens properties."
    )
    doc.add_paragraph("Formula: P_pixel = K * P_camera")
    doc.add_paragraph(
        "The Intrinsic Matrix (K) is derived as:\n"
        "[ fx  s   cx ]\n"
        "[ 0   fy  cy ]\n"
        "[ 0   0   1  ]\n"
        "Derivation:\n"
        "1. Focal Length (fx, fy): Converts meters in camera space into pixels on the sensor.\n"
        "2. Principal Point (cx, cy): Shifts the origin from the lens center to the top-left corner of the image.\n"
        "3. Perspective Division: The 2D coordinates are finally divided by the depth (Z) to simulate how objects "
        "get smaller as they move further away (Perspective Projection)."
    ).paragraph_format.left_indent = Inches(0.3)

    # The 9-Keypoint Model
    doc.add_heading('4. The 9-Keypoint Representation: Robustness & Geometry', level=1)
    doc.add_paragraph(
        "Instead of predicting just the center or the box dimensions, Objectron uses 9 keypoints. "
        "This is a 'Geometric Proxy' representation that the model learns to regress."
    )
    
    doc.add_heading('The Layout:', level=2)
    bullets = [
        "Keypoint 0 (The Anchor): The geometric center of the bottle (Centroid).",
        "Keypoints 1-4 (The Bottom Face): Corners of the box at the base of the bottle.",
        "Keypoints 5-8 (The Top Face): Corners of the box at the top of the bottle."
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    doc.add_heading('Mathematical Rationale:', level=2)
    doc.add_paragraph(
        "Why 9 points instead of a simple 2D box?\n"
        "1. 6-DOF Recovery: With these 9 points, we can solve the PnP (Perspective-n-Point) problem to recover "
        "the exact 3D Rotation and Translation of the bottle from a single image.\n"
        "2. Viewpoint Invariance: Even if part of the bottle is occluded, the model can infer the 'missing' corners "
        "based on the geometric constraints of a cube.\n"
        "3. Shape Consistency: By forcing the model to predict 9 related points, it implicitly learns the "
        "3D structure of the object category (e.g., all bottles are roughly cylindrical/rectangular)."
    ).paragraph_format.left_indent = Inches(0.3)
    
    # Practical Example
    doc.add_heading('5. Practical Interpretation of Annotation Data', level=1)
    doc.add_paragraph(
        "When analyzing a '.pbdata' file, you will encounter the following fields which correspond to the "
        "math above:"
    )
    
    fields = [
        ("scale:", "Corresponds to the 'S' matrix. It defines the physical dimensions (Width, Height, Depth) of the bottle."),
        ("transform:", "A 4x4 matrix representing the View Matrix (V). It changes as the user moves the phone."),
        ("intrinsics:", "A 3x3 matrix (K) representing the camera's lens properties (focal length, principal point)."),
        ("keypoints:", "The final calculated (u, v) normalized coordinates after all transformations are complete.")
    ]
    
    for field, detail in fields:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(field)
        run.font.name = 'Courier New'
        run.bold = True
        p.add_run(f" {detail}")

    # Save
    doc_path = 'Objectron/Objectron_Technical_Deep_Dive.docx'
    doc.save(doc_path)
    print(f"Document saved to {doc_path}")

if __name__ == "__main__":
    create_objectron_doc()
